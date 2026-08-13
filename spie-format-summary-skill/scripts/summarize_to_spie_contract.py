#!/usr/bin/env python
"""Create SPIE_FORMAT_SUMMARY.md and SPIE_INPUT_CONTRACT.md from rough material."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

try:
    from docx import Document
except ImportError:  # pragma: no cover
    Document = None


CANONICAL = [
    "Title",
    "Authors",
    "Affiliations",
    "Corresponding Author",
    "Abstract",
    "Keywords",
    "Sections",
    "Disclosures",
    "Code, Data, and Materials Availability",
    "Acknowledgments",
    "References",
]


ALIASES = {
    "title": "Title",
    "paper title": "Title",
    "authors": "Authors",
    "author": "Authors",
    "affiliations": "Affiliations",
    "affiliation": "Affiliations",
    "corresponding author": "Corresponding Author",
    "correspondence": "Corresponding Author",
    "abstract": "Abstract",
    "keywords": "Keywords",
    "key words": "Keywords",
    "sections": "Sections",
    "body": "Sections",
    "introduction": "Sections",
    "related work": "Sections",
    "methods": "Sections",
    "materials and methods": "Sections",
    "method": "Sections",
    "results": "Sections",
    "discussion": "Sections",
    "conclusion": "Sections",
    "disclosures": "Disclosures",
    "conflicts of interest": "Disclosures",
    "data availability": "Code, Data, and Materials Availability",
    "code, data, and materials availability": "Code, Data, and Materials Availability",
    "acknowledgments": "Acknowledgments",
    "acknowledgements": "Acknowledgments",
    "funding": "Acknowledgments",
    "references": "References",
    "bibliography": "References",
}


BODY_HEADINGS = {
    "introduction",
    "related work",
    "materials and methods",
    "methods",
    "method",
    "methodology",
    "results",
    "discussion",
    "conclusion",
    "conclusions",
}


def read_source(path: Path) -> str:
    if path.suffix.lower() == ".docx":
        if Document is None:
            raise SystemExit("python-docx is required to read DOCX input.")
        doc = Document(path)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [" ".join(cell.text.split()) for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    return path.read_text(encoding="utf-8-sig")


def norm_label(text: str) -> str:
    text = re.sub(r"^#+\s*", "", text.strip())
    text = re.sub(r"[:：]\s*$", "", text)
    text = re.sub(r"^\d+(\.\d+)*\s+", "", text)
    return re.sub(r"\s+", " ", text).lower()


def canonical_heading(line: str) -> str | None:
    label = norm_label(line)
    if label in ALIASES:
        return ALIASES[label]
    if re.match(r"^\d+(\.\d+)*\s+\S+", line.strip()):
        return "Sections"
    return None


def inline_heading(line: str) -> tuple[str | None, str]:
    match = re.match(r"^([A-Za-z][A-Za-z ,/&-]{1,80})[:：]\s*(.+)$", line.strip())
    if not match:
        return None, line
    label = canonical_heading(match.group(1))
    return (label, match.group(2).strip()) if label else (None, line)


def body_heading(line: str) -> str:
    clean = re.sub(r"^#+\s*", "", line.strip())
    label = norm_label(clean)
    if re.match(r"^\d+(\.\d+)*\s+", clean):
        return f"## {clean}"
    number = {
        "introduction": "1",
        "related work": "2",
        "materials and methods": "2",
        "methods": "2",
        "method": "2",
        "methodology": "2",
        "results": "3",
        "discussion": "4",
        "conclusion": "5",
        "conclusions": "5",
    }.get(label)
    if number:
        title = "Materials and Methods" if label in {"method", "methods", "methodology"} else clean.title()
        return f"## {number} {title}"
    return f"## {clean}"


def infer_blocks(text: str) -> dict[str, list[str]]:
    blocks = {name: [] for name in CANONICAL}
    current: str | None = None
    title_seen = False
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            if current:
                blocks[current].append("")
            continue
        inline, rest = inline_heading(line)
        if inline:
            current = inline
            blocks[current].append(rest)
            title_seen = title_seen or current == "Title"
            continue
        heading = canonical_heading(line)
        if heading:
            current = heading
            label = norm_label(line)
            if current == "Sections" and (label in BODY_HEADINGS or re.match(r"^\d+(\.\d+)*\s+\S+", line)):
                blocks[current].append(body_heading(line))
            title_seen = title_seen or current == "Title"
            continue
        if line.startswith("#"):
            heading = canonical_heading(line)
            if heading:
                current = heading
                continue
        if not title_seen:
            blocks["Title"].append(re.sub(r"^#+\s*", "", line))
            title_seen = True
            current = "Authors"
            continue
        if current is None:
            current = "Sections"
        blocks[current].append(line)
    return blocks


def clean(lines: list[str], default: str) -> str:
    text = "\n".join(lines).strip()
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text or default


def numbered_refs(text: str) -> str:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if not lines:
        return "TODO: Add verified numbered references."
    out = []
    for idx, line in enumerate(lines, start=1):
        line = re.sub(r"^\[(\d+)]\s+", r"\1. ", line)
        if not re.match(r"^\d+\.\s+", line):
            line = f"{idx}. {line}"
        out.append(line)
    return "\n".join(out)


def ensure_sections(text: str) -> str:
    defaults = {
        "1": "## 1 Introduction\nTODO: Summarize background, problem gap, contribution, evidence preview, and roadmap.",
        "2": "## 2 Materials and Methods\nTODO: Summarize method, variables, equations, algorithm/workflow, and experimental setup.",
        "3": "## 3 Results\nDATA_NEEDED: Add verified results, figures, tables, metrics, and baselines.",
        "4": "## 4 Discussion\nTODO: Summarize implications, limitations, and scope boundaries.",
        "5": "## 5 Conclusion\nTODO: Summarize contribution and next steps without adding new claims.",
    }
    if not re.search(r"^##\s+\d+\s+", text, flags=re.M):
        return "\n\n".join(defaults.values())
    chunks = re.split(r"(?=^##\s+\d+\s+)", text, flags=re.M)
    by_number: dict[str, str] = {}
    for chunk in chunks:
        chunk = chunk.strip()
        if not chunk:
            continue
        match = re.match(r"^##\s+(\d+)\s+", chunk)
        if match:
            by_number[match.group(1)] = chunk
    for number, default in defaults.items():
        by_number.setdefault(number, default)
    extra = [chunk.strip() for chunk in chunks if chunk.strip() and not re.match(r"^##\s+\d+\s+", chunk.strip())]
    ordered = [by_number[number] for number in sorted(by_number, key=lambda item: int(item) if item.isdigit() else 99)]
    return "\n\n".join([*extra, *ordered])


def keywords(text: str) -> str:
    text = clean(text.splitlines(), "")
    if text:
        return text
    return "TODO: Add 3-6 keywords separated by semicolons."


def inventory(full_text: str) -> dict[str, int]:
    return {
        "references": len(re.findall(r"(^|\n)\s*(\[\d+]|\d+\.)\s+", full_text)),
        "figures": len(re.findall(r"\b(Fig\.|Figure)\s*\d+", full_text, flags=re.I)),
        "tables": len(re.findall(r"\bTable\s*\d+", full_text, flags=re.I)),
        "equations": len(re.findall(r"(\$\$|\\begin\{equation\}|#eq:)", full_text)),
        "citations": len(re.findall(r"\[\d+(?:\s*,\s*\d+)*]", full_text)),
    }


def first_non_todo(text: str, fallback: str) -> str:
    text = " ".join(line.strip() for line in text.splitlines() if line.strip())
    return text if text and not text.startswith("TODO:") else fallback


def build_contract(blocks: dict[str, list[str]]) -> str:
    defaults = {
        "Title": "TODO: Add manuscript title.",
        "Authors": "TODO: Add author names with SPIE affiliation markers, e.g., First Author^1, Second Author^2,*.",
        "Affiliations": "TODO: Add affiliations, one per line.",
        "Corresponding Author": "TODO: Add corresponding author and email.",
        "Abstract": "TODO: Add SPIE abstract: what, why hard, how, evidence, strongest result or evidence boundary.",
        "Keywords": "TODO: Add 3-6 keywords separated by semicolons.",
        "Sections": "## 1 Introduction\nTODO: Add manuscript body.",
        "Disclosures": "The authors declare that there are no conflicts of interest.",
        "Code, Data, and Materials Availability": "DATA_NEEDED: Add code, data, and materials availability statement.",
        "Acknowledgments": "TODO: Add acknowledgments and funding statement.",
        "References": "TODO: Add verified numbered references.",
    }
    parts = []
    for name in CANONICAL:
        content = clean(blocks[name], defaults[name])
        if name == "Sections":
            content = ensure_sections(content)
        elif name == "References":
            content = numbered_refs(content)
        elif name == "Keywords":
            content = keywords(content)
        parts.append(f"# {name}\n{content}")
    return "\n\n".join(parts) + "\n"


def build_summary(blocks: dict[str, list[str]], full_text: str, contract_path: Path) -> str:
    inv = inventory(full_text)
    title = first_non_todo(clean(blocks["Title"], ""), "TODO: Add manuscript title.")
    abstract = first_non_todo(clean(blocks["Abstract"], ""), "TODO: Add abstract.")
    normalized_sections = ensure_sections(clean(blocks["Sections"], ""))
    numbered_sections = "present" if re.search(r"^##\s+\d+\s+", normalized_sections, flags=re.M) else "needs normalization"
    return "\n".join(
        [
            "# SPIE_FORMAT_SUMMARY",
            "",
            "## Target",
            "- Journal/format: SPIE journal manuscript",
            "- Output expected: SPIE input contract or DOCX via downstream SPIE paper agent",
            "",
            "## Paper Identity",
            f"- Proposed title: {title}",
            "- Research topic: TODO: Confirm concise topic.",
            f"- Authors: {first_non_todo(clean(blocks['Authors'], ''), 'TODO: Add authors.')}",
            f"- Affiliations: {first_non_todo(clean(blocks['Affiliations'], ''), 'TODO: Add affiliations.')}",
            f"- Corresponding author: {first_non_todo(clean(blocks['Corresponding Author'], ''), 'TODO: Add corresponding author.')}",
            "",
            "## Research Story",
            f"- Background/problem summary: {abstract}",
            "- Problem/gap: TODO: Extract or add the precise research gap.",
            "- Method: TODO: Extract method, algorithm, model, or workflow.",
            "- Data/experiment: DATA_NEEDED: Add verified data, baselines, metrics, and settings.",
            "- Main result: DATA_NEEDED: Add strongest verified result.",
            "- Contribution: TODO: Add one-sentence contribution.",
            "- Limitations: TODO: Add limitations and scope boundaries.",
            "",
            "## Evidence Inventory",
            f"- References detected: {inv['references']}",
            f"- Citation markers detected: {inv['citations']}",
            f"- Figures detected: {inv['figures']}",
            f"- Tables detected: {inv['tables']}",
            f"- Equations detected: {inv['equations']}",
            "- Raw data: DATA_NEEDED unless supplied separately.",
            "",
            "## SPIE Format Readiness",
            f"- Abstract: {'present' if clean(blocks['Abstract'], '') else 'missing'}",
            f"- Keywords: {'present' if clean(blocks['Keywords'], '') else 'missing'}",
            f"- Numbered sections: {numbered_sections}",
            f"- Disclosures: {'present' if clean(blocks['Disclosures'], '') else 'default inserted'}",
            f"- Data availability: {'present' if clean(blocks['Code, Data, and Materials Availability'], '') else 'missing'}",
            f"- Acknowledgments: {'present' if clean(blocks['Acknowledgments'], '') else 'missing'}",
            f"- References: {'present' if clean(blocks['References'], '') else 'missing'}",
            "",
            "## Handoff Notes",
            "- Safe claims: Use only claims explicitly supported by supplied text, citations, tables, figures, or data.",
            "- Claims needing evidence: Any numerical, comparative, or performance claim without raw data should be marked `DATA_NEEDED:`.",
            "- Data needed from user: result tables, figure files, captions, equations, baselines, metrics, DOI/full references, author metadata.",
            "",
            f"Next agent input: `{contract_path}`",
            "",
        ]
    )


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("source", type=Path, help="Input material: .md, .txt, or .docx.")
    parser.add_argument("--out-dir", type=Path, required=True, help="Output directory.")
    parser.add_argument("--summary-name", default="SPIE_FORMAT_SUMMARY.md")
    parser.add_argument("--contract-name", default="SPIE_INPUT_CONTRACT.md")
    args = parser.parse_args()

    text = read_source(args.source)
    blocks = infer_blocks(text)
    args.out_dir.mkdir(parents=True, exist_ok=True)
    contract_path = args.out_dir / args.contract_name
    summary_path = args.out_dir / args.summary_name
    contract_path.write_text(build_contract(blocks), encoding="utf-8")
    summary_path.write_text(build_summary(blocks, text, contract_path), encoding="utf-8")
    print(f"Summary: {summary_path}")
    print(f"Contract: {contract_path}")


if __name__ == "__main__":
    main()
